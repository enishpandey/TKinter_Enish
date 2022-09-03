import tkinter
from tkinter import *
import docx
from docx import Document

from docx.opc.coreprops import CoreProperties
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Cm, Pt
from tkinter import messagebox
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()
root = tkinter.Tk()

root['bg']="white"
root.title("Question Generator")


top = Frame(root, bg="white")
top.grid(row=0, column=1)

top1 = Frame(root, bg="white")
top1.grid(row=0, column=0)



# font for the question heading and questions
h = ('calibri', 20, 'bold')
oh = ('calibri', 10, 'bold')
oi = ('calibri', 10)

# to clear all the inputs
def clear_all():
    # head_title.destroy()
    very_short.destroy()
    short.destroy()
    fill_in_the_blanks.destroy()
    true_false.destroy()
    full_form.destroy()
    match_the_following.destroy()
    number_system.destroy()
    technical_terms.destroy()
    short_notes.destroy()


def generate():
    tn=terminal_name_entry.get()
    sn=school_name_entry.get()
    sa=school_address_entry.get()

    terminal_name_text=f'''{tn}'''

    school_name_text=f'''{sn}'''

    school_address_text=f'''{sa}'''


    doc=Document()
    sections=doc.sections

    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    section = doc.sections[0]






    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before=Pt(5)
    para.paragraph_format.space_after=Pt(0)
    # para.paragraph_format.line_spacing=WD_LINE_SPACING.MULTIPLE
    paragra = para.add_run(f'{terminal_name_text}')
    paragra.font.name = 'Calibri'
    paragra.font.size = Pt(16)
    paragra.font.bold=True

    # for _ in range(1):
    #     linespace_style = doc.styles['Body Text']
    #     linespace = doc.add_paragraph(style=linespace_style).add_run()
    #     linespace.font.size = 80


    # school_name_style = doc.styles['Body Text']
    # school_name_line = doc.add_paragraph(style=school_name_style).add_run()
    # school_name_line.font.size = Pt(20)
    # school_name_line.font.bold=True

    school_name_line = doc.add_paragraph()
    school_name_line.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # para.paragraph_format.space_before=Pt(0)
    # para.paragraph_format.space_after=Pt(0)
    paragra1 = school_name_line.add_run(f'{school_name_text}')
    paragra1.font.name = 'Calibri'
    paragra1.font.size = Pt(20)
    paragra1.font.bold=True


    school_address_line = doc.add_paragraph()
    school_address_line.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # para.paragraph_format.space_before=Pt(0)
    # para.paragraph_format.space_after=Pt(0)
    address1 = school_address_line.add_run(f'{school_address_text}')
    address1.font.size = Pt(11)
    address1.font.name = 'Calibri'
    address1.font.bold=True




    doc.save("question1.docx")
# head_title = Label(root,text="Questions for Class-name", font=("normal", 15, "bold"), bg="white")
# head_title.grid()
#Labels

terminal_name = Label(top1, text="Terminal Name:",font=("normal", 10, "bold"), bg="white")
terminal_name.grid(row=2 , column=0, sticky = W, padx=20, pady=15)

school_name = Label(top1, text="School Name",font=("normal", 10, "bold"), bg="white")
school_name.grid(row=3 , column=0, sticky = W, padx=20, pady=15)


school_address = Label(top1, text="School Address:",font=("normal", 10, "bold"), bg="white")
school_address.grid(row=4, column=0, sticky = W, padx=20, pady=15)



# Entry


terminal_name_entry = Entry(top1, width=30, bg="#EAFAFA", fg="black")
terminal_name_entry.grid(row=2, column=1)

school_name_entry = Entry(top1, width=30, bg="#EAFAFA", fg="black")
school_name_entry.grid(row=3, column=1)

school_address_entry = Entry(top1, width=30, bg="#EAFAFA", fg="black")
school_address_entry.grid(row=4, column=1)



# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------




#Labels
very_short = Label(top, text="Very Short:",font=("normal", 10, "bold"), bg="white")
very_short.grid(row=1 , column=0, sticky = W, padx=20, pady=15)


short = Label(top, text="Short Questions:",font=("normal", 10, "bold"), bg="white")
short.grid(row=2 , column=0, sticky = W, padx=20, pady=15)

fill_in_the_blanks = Label(top, text="Fill in the blanks:",font=("normal", 10, "bold"), bg="white")
fill_in_the_blanks.grid(row=3 , column=0, sticky = W, padx=20, pady=15)


true_false = Label(top, text="True or False:",font=("normal", 10, "bold"), bg="white")
true_false.grid(row=4, column=0, sticky = W, padx=20, pady=15)

full_form = Label(top, text="Full Form:",font=("normal", 10, "bold"), bg="white")
full_form.grid(row=5 , column=0, sticky = W, padx=20, pady=15)

match_the_following = Label(top, text="Match the following:",font=("normal", 10, "bold"), bg="white")
match_the_following.grid(row=6 , column=0, sticky = W, padx=20, pady=15)

number_system = Label(top, text="Number System:",font=("normal", 10, "bold"), bg="white")
number_system.grid(row=7 , column=0, sticky = W, padx=20, pady=15)

technical_terms = Label(top, text="Technical Terms:",font=("normal", 10, "bold"), bg="white")
technical_terms.grid(row=8 , column=0, sticky = W, padx=20, pady=15)

short_notes = Label(top, text="Short Notes:",font=("normal", 10, "bold"), bg="white")
short_notes.grid(row=9 , column=0, sticky = W, padx=20, pady=15)

# Entry
very_short_Button = Button(top,text="Click to get VSQ", width=30, bg="lightblue", fg="black")
very_short_Button.grid(row=1, column=1)

short_Button = Button(top,text="Click to get SQ", width=30, bg="lightblue", fg="black")
short_Button.grid(row=2, column=1)

fill_in_the_blanks_Button = Button(top,text="Click to get FITB", width=30, bg="lightblue", fg="black")
fill_in_the_blanks_Button.grid(row=3, column=1)

true_false_Button = Button(top,text="Click to get TF", width=30, bg="lightblue", fg="black")
true_false_Button.grid(row=4, column=1)

full_form_Button = Button(top,text="Click to get FF", width=30, bg="lightblue", fg="black")
full_form_Button.grid(row=5, column=1)

match_the_following_Button = Button(top,text="Click to get MTF", width=30, bg="lightblue", fg="black")
match_the_following_Button.grid(row=6, column=1)

number_system_Button = Button(top,text="Click to get NS", width=30, bg="lightblue", fg="black")
number_system_Button.grid(row=7, column=1)

technical_terms_Button = Button(top,text="Click to get TT", width=30, bg="lightblue", fg="black")
technical_terms_Button.grid(row=8, column=1)

short_notes_Button = Button(top,text="Click to get SN", width=30, bg="lightblue", fg="black")
short_notes_Button.grid(row=9, column=1)


generate_doc = Button(root,text="Generate word", width=15,height="2", bg="green", fg="white", command=generate)
generate_doc.grid(row=10, column=0)

clear = Button(root,text="Clear", width=15,height="2", bg="green", fg="white", command=clear_all)
clear.grid(row=10, column=1)

exit = Button(root,text="Exit", width=15,height="2", bg="green", fg="white")
exit.grid(row=10, column=2)

root.geometry("980x800")
root.mainloop()

# head = doc.add_heading("Welcome to Nepal", level=2)
# doc.add_page_break()
# head2 = doc.add_heading("Welcome to Nepal", level=2)