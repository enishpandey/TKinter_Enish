import tkinter
from tkinter import *
from docx import Document
from docx.opc.coreprops import CoreProperties
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Cm, Pt
from tkinter import messagebox


top = tkinter.Tk()

top['bg']="white"
top.title("Resignation Letter Generator")

# font for the question heading and questions
h = ('calibri', 20, 'bold')
oh = ('calibri', 10, 'bold')
oi = ('calibri', 10)

# to clear all the inputs
def clear_inputs():
    date_of_letter_entry.delete(0,'end')
    supervisor_name_entry.delete(0,'end')
    supervisor_position_entry.delete(0,'end')
    address_entry.delete(0,'end')
    position_held_entry.delete(0,'end')
    resignation_date_requested_entry.delete(0,'end')
    reason_for_resignation_entry.delete(0,'end')
    resignee_name_entry.delete(0,'end')
    address_of_resignee_entry.delete(0,'end')
    company_name_entry.delete(0,'end')
def generate_doc():

    a=date_of_letter_entry.get()
    b=supervisor_name_entry.get()
    c=supervisor_position_entry.get()
    d=address_entry.get()
    e=position_held_entry.get()
    m=resignation_date_requested_entry.get()
    g=reason_for_resignation_entry.get()
    h=resignee_name_entry.get()
    i=address_of_resignee_entry.get()
    j=company_name_entry.get()

    head_text=f'''
{a}

{b}
{c}
{d}
'''
    body_text=f'''
Dear{b},

I am writing to formally inform you of my resignation from my position as {e} at {j}. In accordance with the period of notice agreed within my contract, my last day will be {m}.
   
{g}. Therefore, I would be grateful if you could confirm receipt of my notification and of my leaving date.
   
I would like to take this opportunity to thank you for all of the opportunities presented to me within the period of my employment. I have enjoyed my time working at {j}, however, in the best interests of my career, I feel that the time is right to move on.

Finally, if there is anything I can do to ensure a smooth and efficient handover process, please do not hesitate to let me know. 

I wish you all the very best for the future.

Thank you.

Yours sincerely,
{h}
{i}
    '''

    doc=Document()
    sections = doc.sections

    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    section = doc.sections[0]

#     Headings
    head_style = doc.styles['Body Text']
    head = doc.add_paragraph(style=head_style).add_run(f'{head_text}')
    head.font.size = Pt(12)
    head.font.name = 'Times New Roman'

#     Body Text
    body_style = doc.styles['Body Text']
    body = doc.add_paragraph(style=body_style).add_run(f'{body_text}')
    body.font.size = Pt(14)
    body.font.name='Times New Roman'

    doc.save("Resignation.docx")



# head_title = Label(top,text="Questions for Class-name", font=("normal", 15, "bold"), bg="white")
# head_title.grid(row=0 , column=1)

    # date_of_letter_entry
    # supervisor_name_entry
    # supervisor_position_entry
    # address_entry
    # position_held_entry
    # resignation_date_requested_entry
    # reason_for_resignation_entry
    # resignee_name_entry
    # address_of_resignee_entry

#Labels
# day = Label(top, text="Enter a Day:",font=("normal", 10, "bold"), bg="white")
# day.grid(row=1 , column=0, sticky = W, padx=20, pady=15)
#
#
# month = Label(top, text="Enter a Month:",font=("normal", 10, "bold"), bg="white")
# month.grid(row=2 , column=0, sticky = W, padx=20, pady=15)

date_of_letter = Label(top, text="Enter a Date:",font=("normal", 10, "bold"), bg="white")
date_of_letter.grid(row=3 , column=0, sticky = W, padx=20, pady=15)

supervisor_name = Label(top, text="Enter Supervisor Name:",font=("normal", 10, "bold"), bg="white")
supervisor_name.grid(row=4, column=0, sticky = W, padx=20, pady=15)

supervisor_position = Label(top, text="Supervisor Position:",font=("normal", 10, "bold"), bg="white")
supervisor_position.grid(row=5 , column=0, sticky = W, padx=20, pady=15)

address = Label(top, text="Enter Address:",font=("normal", 10, "bold"), bg="white")
address.grid(row=6 , column=0, sticky = W, padx=20, pady=15)

position_held = Label(top, text="Position Held:",font=("normal", 10, "bold"), bg="white")
position_held.grid(row=7 , column=0, sticky = W, padx=20, pady=15)

resignation_date_requested = Label(top, text="Resignation Date Requested:",font=("normal", 10, "bold"), bg="white")
resignation_date_requested.grid(row=8 , column=0, sticky = W, padx=20, pady=15)

reason_for_resignation = Label(top, text="Reason for Resignation:",font=("normal", 10, "bold"), bg="white")
reason_for_resignation.grid(row=9 , column=0, sticky = W, padx=20, pady=15)

resignee_name = Label(top, text="Resignee Name:",font=("normal", 10, "bold"), bg="white")
resignee_name.grid(row=10 , column=0, sticky = W, padx=20, pady=15)

address_of_resignee = Label(top, text="Address of Resignee:",font=("normal", 10, "bold"), bg="white")
address_of_resignee.grid(row=11 , column=0, sticky = W, padx=20, pady=15)

company_name= Label(top, text="Company Name:",font=("normal", 10, "bold"), bg="white")
company_name.grid(row=12 , column=0, sticky = W, padx=20, pady=15)


# Entry
# day_entry = Entry(top,width=10, fg="black", bd=3)
# day_entry.grid(row=1, column=1)
#
# month_entry = Entry(top,width=10, fg="black", bd=3)
# month_entry.grid(row=2, column=1)

date_of_letter_entry = Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
date_of_letter_entry.grid(row=3, column=1, sticky=W)

supervisor_name_entry = Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
supervisor_name_entry.grid(row=4, column=1, sticky=W)

supervisor_position_entry= Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
supervisor_position_entry.grid(row=5, column=1, sticky=W)

address_entry= Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
address_entry.grid(row=6, column=1, sticky=W)

position_held_entry= Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
position_held_entry.grid(row=7, column=1, sticky=W)

resignation_date_requested_entry= Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
resignation_date_requested_entry.grid(row=8, column=1, sticky=W)

reason_for_resignation_entry= Entry(top,width=50, fg="black", bd=3, bg="#EAFAFA")
reason_for_resignation_entry.grid(row=9, column=1, sticky=W)

resignee_name_entry= Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
resignee_name_entry.grid(row=10, column=1, sticky=W)

address_of_resignee_entry= Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
address_of_resignee_entry.grid(row=11, column=1, sticky=W)

company_name_entry= Entry(top,width=30, fg="black", bd=3, bg="#EAFAFA")
company_name_entry.grid(row=12, column=1, sticky=W)


# very_short_Button = Button(top,text="Click to get VSQ", width=30, bg="lightblue", fg="black")
# very_short_Button.grid(row=1, column=1)
#
# short_Button = Button(top,text="Click to get SQ", width=30, bg="lightblue", fg="black")
# short_Button.grid(row=2, column=1)
#
# fill_in_the_blanks_Button = Button(top,text="Click to get FITB", width=30, bg="lightblue", fg="black")
# fill_in_the_blanks_Button.grid(row=3, column=1)
#
# true_false_Button = Button(top,text="Click to get TF", width=30, bg="lightblue", fg="black")
# true_false_Button.grid(row=4, column=1)
#
# full_form_Button = Button(top,text="Click to get FF", width=30, bg="lightblue", fg="black")
# full_form_Button.grid(row=5, column=1)
#
# match_the_following_Button = Button(top,text="Click to get MTF", width=30, bg="lightblue", fg="black")
# match_the_following_Button.grid(row=6, column=1)
#
# number_system_Button = Button(top,text="Click to get NS", width=30, bg="lightblue", fg="black")
# number_system_Button.grid(row=7, column=1)
#
# technical_terms_Button = Button(top,text="Click to get TT", width=30, bg="lightblue", fg="black")
# technical_terms_Button.grid(row=8, column=1)
#
# short_notes_Button = Button(top,text="Click to get SN", width=30, bg="lightblue", fg="black")
# short_notes_Button.grid(row=9, column=1)

generate = Button(top,text="Generate word", width=15,height="2", bg="green", fg="white", command=generate_doc)
generate.grid(row=13, column=0)

clear = Button(top,text="Clear", width=15,height="2", bg="green", fg="white", command=clear_inputs)
clear.grid(row=13, column=1)



top.geometry(("650x620"))
top.mainloop()